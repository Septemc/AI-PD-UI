from docx import Document
import paddlehub as hub
import torch
import numpy as np
import jieba
import json
from models.Nets import TextCNN
from utils.DataProcess import read_and_segment_txt


def txt_detection(file_path, model_path):
    vocab_size = 800
    embedding_dim = 300
    num_filters = 128
    output_dim = 2
    target_shape = (vocab_size, 300)

    # 加载模型文件
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model = TextCNN(embedding_dim, num_filters, output_dim)
    model.load_state_dict(torch.load(model_path, map_location=device))

    # 调用函数进行读取和分段处理
    paragraphs = read_and_segment_txt(file_path)

    text_chunks = []
    word_num = 0

    for paragraph in paragraphs:
        text = paragraph
        if text:
            text_chunks.append(text)
            word_num += len(text)

    embedding = hub.Module(name='w2v_wiki_target_word-word_dim300')
    data_text_tensor = []

    for text in text_chunks:
        words = jieba.cut(text)
        words_list = list(words)
        text_vector = embedding.search(words_list)
        text_vector_np = np.array(text_vector)
        pad_rows = target_shape[0] - text_vector_np.shape[0]
        padding = np.zeros((pad_rows, 300))
        text_padded_data = np.vstack([text_vector_np, padding])
        text_vector_tensor = torch.tensor(text_padded_data).float()
        data_text_tensor.append(text_vector_tensor)

    text_datas = torch.stack(data_text_tensor)
    outputs = model(text_datas)

    total_weighted_ai_prob = 0
    total_weighted_real_prob = 0
    result_details = []

    for i, text in enumerate(text_chunks):
        text_length = len(text)
        weighted_ai_prob = outputs[i][0].item() * 100 * (text_length / word_num)
        weighted_real_prob = outputs[i][1].item() * 100 * (text_length / word_num)
        total_weighted_ai_prob += weighted_ai_prob
        total_weighted_real_prob += weighted_real_prob

        result_details.append({
            "text_number": i + 1,
            "text": text,
            "AI_probability": round(outputs[i][0].item() * 100, 3),  # 保留三位小数
            "real_probability": round(outputs[i][1].item() * 100, 3)  # 保留三位小数
        })

    # 计算总概率时保留三位小数
    total_weighted_ai_prob = round(total_weighted_ai_prob, 3)
    total_weighted_real_prob = round(total_weighted_real_prob, 3)

    result = {
        "whole_ai_prob": total_weighted_ai_prob,
        "whole_real_prob": total_weighted_real_prob,
        "details": result_details
    }

    # 将Python字典转换为JSON格式的字符串
    result_json = json.dumps(result, ensure_ascii=False, indent=4)

    return result_json


def doc_detection(file_path, model_path):
    vocab_size = 800
    embedding_dim = 300
    num_filters = 128
    output_dim = 2
    target_shape = (vocab_size, 300)


    # 加载模型文件
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model = TextCNN(embedding_dim, num_filters, output_dim)
    model.load_state_dict(torch.load(model_path, map_location=device))

    doc = Document(file_path)
    text_chunks = []
    word_num = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text:
            text_chunks.append(text)
            word_num += len(text)

    embedding = hub.Module(name='w2v_wiki_target_word-word_dim300')
    data_text_tensor = []

    for text in text_chunks:
        words = jieba.cut(text)
        words_list = list(words)
        text_vector = embedding.search(words_list)
        text_vector_np = np.array(text_vector)
        pad_rows = target_shape[0] - text_vector_np.shape[0]
        padding = np.zeros((pad_rows, 300))
        text_padded_data = np.vstack([text_vector_np, padding])
        text_vector_tensor = torch.tensor(text_padded_data).float()
        data_text_tensor.append(text_vector_tensor)

    text_datas = torch.stack(data_text_tensor)
    outputs = model(text_datas)

    total_weighted_ai_prob = 0
    total_weighted_real_prob = 0
    result_details = []

    for i, text in enumerate(text_chunks):
        text_length = len(text)
        weighted_ai_prob = outputs[i][0].item() * 100 * (text_length / word_num)
        weighted_real_prob = outputs[i][1].item() * 100 * (text_length / word_num)
        total_weighted_ai_prob += weighted_ai_prob
        total_weighted_real_prob += weighted_real_prob

        result_details.append({
            "text_number": i + 1,
            "text": text,
            "AI_probability": round(outputs[i][0].item() * 100, 3),  # 保留三位小数
            "real_probability": round(outputs[i][1].item() * 100, 3)  # 保留三位小数
        })

    # 计算总概率时保留三位小数
    total_weighted_ai_prob = round(total_weighted_ai_prob, 3)
    total_weighted_real_prob = round(total_weighted_real_prob, 3)

    result = {
        "whole_ai_prob": total_weighted_ai_prob,
        "whole_real_prob": total_weighted_real_prob,
        "details": result_details
    }

    # 将Python字典转换为JSON格式的字符串
    result_json = json.dumps(result, ensure_ascii=False, indent=4)

    return result_json


def pdf_detection(file_path, model_path):
    vocab_size = 800
    embedding_dim = 300
    num_filters = 128
    output_dim = 2
    target_shape = (vocab_size, 300)

    # 加载模型文件
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model = TextCNN(embedding_dim, num_filters, output_dim)
    model.load_state_dict(torch.load(model_path, map_location=device))

    doc = Document(file_path)
    text_chunks = []
    word_num = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text:
            text_chunks.append(text)
            word_num += len(text)

    embedding = hub.Module(name='w2v_wiki_target_word-word_dim300')
    data_text_tensor = []

    for text in text_chunks:
        words = jieba.cut(text)
        words_list = list(words)
        text_vector = embedding.search(words_list)
        text_vector_np = np.array(text_vector)
        pad_rows = target_shape[0] - text_vector_np.shape[0]
        padding = np.zeros((pad_rows, 300))
        text_padded_data = np.vstack([text_vector_np, padding])
        text_vector_tensor = torch.tensor(text_padded_data).float()
        data_text_tensor.append(text_vector_tensor)

    text_datas = torch.stack(data_text_tensor)
    outputs = model(text_datas)

    total_weighted_ai_prob = 0
    total_weighted_real_prob = 0
    result_details = []

    for i, text in enumerate(text_chunks):
        text_length = len(text)
        weighted_ai_prob = outputs[i][0].item() * 100 * (text_length / word_num)
        weighted_real_prob = outputs[i][1].item() * 100 * (text_length / word_num)
        total_weighted_ai_prob += weighted_ai_prob
        total_weighted_real_prob += weighted_real_prob

        result_details.append({
            "text_number": i + 1,
            "text": text,
            "AI_probability": round(outputs[i][0].item() * 100, 3),  # 保留三位小数
            "real_probability": round(outputs[i][1].item() * 100, 3)  # 保留三位小数
        })

    # 计算总概率时保留三位小数
    total_weighted_ai_prob = round(total_weighted_ai_prob, 3)
    total_weighted_real_prob = round(total_weighted_real_prob, 3)

    result = {
        "whole_ai_prob": total_weighted_ai_prob,
        "whole_real_prob": total_weighted_real_prob,
        "details": result_details
    }

    # 将Python字典转换为JSON格式的字符串
    result_json = json.dumps(result, ensure_ascii=False, indent=4)

    return result_json
